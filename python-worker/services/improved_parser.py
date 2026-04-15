"""
Improved Resume Parser with Better Extraction Logic
"""
import re
from config import settings
from typing import Dict, Any, List, Optional
from pathlib import Path
from loguru import logger
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
import json

class ImprovedResumeParser:
    """Enhanced parser with better entity extraction"""
    
    def __init__(self):
        self.nlp = None
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        # Updated phone pattern to handle Indian format with spaces: +91 95885 30839
        # Matches: +91 9588530839, +91-9588530839, +91 95885 30839, 9588530839
        self.phone_pattern = re.compile(r'(?:\+91[\s-]*)?([6-9]\d[\s-]*\d{3}[\s-]*\d{5})')
        self.linkedin_pattern = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/([\w-]+)', re.I)
        self.github_pattern = re.compile(r'(?:https?://)?(?:www\.)?github\.com/([\w-]+)', re.I)
        
        # Extended skills database (500+ skills)
        self.skills_database = self._load_skills_database()

    # ------------------------------------------------------------------
    # Internal path helper
    # ------------------------------------------------------------------
    def _ensure_abs_under_storage(self, path_str: str) -> str:
        """
        Normalize to absolute path WITHOUT creating storage/storage/... bugs.
        """
        p = Path(path_str)

        # If already absolute, return as-is
        if p.is_absolute():
            return str(p.resolve())

        # Strip leading "storage/" from relative paths
        parts = p.parts
        if parts and parts[0].lower() == "storage":
            p = Path(*parts[1:])

        # Join with storage root
        final_path = (settings.storage_path / p).resolve()
        return str(final_path)
        
    def _load_skills_database(self) -> set:
        """Load comprehensive skills database"""
        skills = {
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
            'go', 'rust', 'scala', 'r', 'matlab', 'perl', 'bash', 'shell', 'powershell',
            
            # Web Frameworks
            'react', 'angular', 'vue', 'vue.js', 'react.js', 'next.js', 'nuxt.js', 'svelte',
            'node.js', 'express', 'express.js', 'django', 'flask', 'fastapi', 'spring', 'spring boot',
            'asp.net', '.net', '.net core', 'laravel', 'symfony', 'ruby on rails',
            
            # Databases
            'sql', 'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'dynamodb', 'oracle', 'sql server', 'mariadb', 'sqlite', 'neo4j', 'firebase',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'k8s', 'terraform',
            'ansible', 'jenkins', 'gitlab ci', 'github actions', 'circleci', 'travis ci',
            'ci/cd', 'devops', 'cloudformation', 'helm', 'rancher', 'openshift',
            
            # AI/ML/Data Science
            'machine learning', 'deep learning', 'nlp', 'natural language processing',
            'computer vision', 'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas',
            'numpy', 'opencv', 'spacy', 'hugging face', 'transformers', 'bert', 'gpt',
            'data science', 'data analysis', 'data engineering', 'big data', 'spark', 'hadoop',
            
            # Mobile Development
            'android', 'ios', 'react native', 'flutter', 'swift', 'swiftui', 'kotlin',
            'xamarin', 'ionic', 'cordova',
            
            # Tools & Technologies
            'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
            'vs code', 'visual studio', 'intellij', 'eclipse', 'postman', 'swagger',
            
            # Testing
            'jest', 'mocha', 'chai', 'pytest', 'unittest', 'selenium', 'cypress',
            'junit', 'testng', 'jasmine', 'karma',
            
            # Methodologies
            'agile', 'scrum', 'kanban', 'waterfall', 'tdd', 'bdd', 'microservices',
            'rest api', 'restful', 'graphql', 'grpc', 'soap', 'websocket',
            
            # Other
            'html', 'css', 'sass', 'less', 'tailwind', 'bootstrap', 'material ui',
            'redux', 'vuex', 'webpack', 'vite', 'babel', 'eslint', 'prettier',
            'nginx', 'apache', 'linux', 'ubuntu', 'windows', 'macos'
        }
        return skills
        
    async def load_models(self):
        """Load models (Skipped: spaCy removed to save space)"""
        self.nlp = None
    
    async def parse_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse resume file and extract structured information"""
        try:
            logger.info(f"📄 Parsing file: {filename}")
            
            # Extract text based on file type
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                raw_text = await self._extract_from_pdf(self._ensure_abs_under_storage(file_path))
            elif file_ext == '.docx':
                raw_text = await self._extract_from_docx(self._ensure_abs_under_storage(file_path))
            elif file_ext == '.doc':
                raw_text = await self._extract_from_doc(self._ensure_abs_under_storage(file_path))
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("Unable to extract meaningful text from resume")
            
            # Extract structured information
            parsed_data = await self._extract_entities(raw_text)
            
            # Add metadata
            parsed_data['raw_text'] = raw_text[:10000]  # Limit to 10k chars
            parsed_data['filename'] = filename
            parsed_data['file_type'] = file_ext
            parsed_data['parsed_at'] = datetime.utcnow().isoformat()
            parsed_data['text_length'] = len(raw_text)
            
            logger.info(f"✅ Successfully parsed {filename}")
            logger.info(f"   - Name: {parsed_data.get('full_name')}")
            logger.info(f"   - Email: {parsed_data.get('email')}")
            logger.info(f"   - Skills: {len(parsed_data.get('skills', []))} found")
            logger.info(f"   - Experience: {parsed_data.get('total_experience_years', 0)} years")
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF for better spacing"""
        # Always normalize path under storage to avoid storage/storage/... bugs
        abs_path = self._ensure_abs_under_storage(file_path)
        try:
            doc = fitz.open(abs_path)
            text_parts = []
            
            for page in doc:
                # Extract text with proper spacing
                text = page.get_text("text")
                text_parts.append(text)
            
            doc.close()
            
            full_text = "\n".join(text_parts).strip()
            
            if full_text and len(full_text.strip()) > 50:
                logger.info(f"✅ PDF text extracted successfully using direct text layer ({len(full_text)} chars)")
                return full_text
            
            # If text extraction returned empty, use OCR fallback
            logger.warning("PDF text extraction returned empty/insufficient text, switching to OCR fallback...")
            return await self._extract_with_ocr(abs_path)
            
        except Exception as e:
            logger.warning(f"PDF extraction encountered issue ({e}), using OCR fallback...")
            return await self._extract_with_ocr(abs_path)
    
    async def _extract_with_ocr(self, file_path: str) -> str:
        """Extract text using OCR fallback for scanned/image PDFs"""
        abs_path = self._ensure_abs_under_storage(file_path)
        try:
            from services.ocr_parser import OCRParser
            ocr_parser = OCRParser()
            logger.info("🔄 Attempting OCR extraction for image-based PDF...")
            ocr_text = await ocr_parser.extract_text(abs_path)
            
            if ocr_text and len(ocr_text.strip()) > 50:
                logger.info(f"✅ OCR extraction successful: {len(ocr_text)} characters")
                return ocr_text
            else:
                raise ValueError("OCR extraction returned empty text")
                
        except Exception as ocr_error:
            logger.error(f"❌ OCR fallback failed: {ocr_error}")
            raise ValueError(f"Both PDF text extraction and OCR fallback failed: {ocr_error}")
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
            
    async def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from legacy .doc files using Spire.Doc or Fallbacks"""
        abs_path = self._ensure_abs_under_storage(file_path)
        try:
            # First try Spire.Doc if installed
            try:
                from spire.doc import Document as SpireDocument
                doc = SpireDocument()
                doc.LoadFromFile(abs_path)
                text = doc.GetText()
                # Remove evaluation warning common in Spire.Doc
                text = text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "")
                if text and text.strip():
                    return text.strip()
            except ImportError:
                logger.warning("Spire.Doc not installed. Attempting fallbacks for .doc")
            except Exception as e:
                logger.warning(f"Spire.Doc extraction failed: {e}")

            # Try antiword via subprocess (common on Linux)
            import subprocess
            try:
                result = subprocess.run(['antiword', abs_path], capture_output=True, text=True, check=True)
                if result.stdout:
                    return result.stdout.strip()
            except (subprocess.SubprocessError, FileNotFoundError):
                pass
                
            raise ValueError("No suitable tool found to parse .doc file. Please install Spire.Doc or antiword.")
            
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            raise
    
    async def _extract_entities(self, text: str) -> Dict[str, Any]:
        """Extract structured entities from resume text"""
        result = {
            # Personal Details
            'email': None,
            'phone': None,
            'full_name': None,
            'location': None,
            
            # Required Details
            'gender': None,
            'designation': None,
            'total_experience': 0.0,
            'deployment_type': None,
            'availability': None,
            'country': None,
            'city': None,
            'primary_skills': [],
            'secondary_skills': [],
            
            # Additional Details
            'experience': [],  # [{company, role, start_year, end_year}]
            'projects': [],    # [{project_title, description}]
            'education': [],   # [{degree, institute, passing_year}]
            
            # Legacy fields (for compatibility)
            'linkedin_url': None,
            'github_url': None,
            'skills': []  # Will be split into primary/secondary
        }
        
        # Extract contact information
        result['email'] = self._extract_email(text)
        result['phone'] = self._extract_phone(text)
        result['linkedin_url'] = self._extract_linkedin(text)
        result['github_url'] = self._extract_github(text)
        
        # Use spaCy for NER
        if self.nlp:
            doc = self.nlp(text[:5000])  # Process first 5000 chars for speed
            
            # Extract name (better logic)
            result['full_name'] = self._extract_name(text, doc)
            
            # Extract organizations and use first as current company
            orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
            if orgs:
                result['current_company'] = orgs[0]
            
            # Extract location
            locations = [ent.text for ent in doc.ents if ent.label_ in ("GPE", "LOC")]
            if locations:
                result['location'] = locations[0]
        
        # Extract skills and split into primary/secondary
        all_skills = self._extract_skills_improved(text)
        result['skills'] = all_skills
        result['primary_skills'] = all_skills[:5] if len(all_skills) > 5 else all_skills
        result['secondary_skills'] = all_skills[5:] if len(all_skills) > 5 else []
        
        # Extract designation
        result['designation'] = self._extract_designation(text)
        
        # Extract location details
        location_data = self._extract_location_details(text)
        result['country'] = location_data.get('country')
        result['city'] = location_data.get('city')
        
        # Extract experience (new format: company, role, start_year, end_year)
        result['experience'] = self._extract_experience_new_format(text)
        result['total_experience'] = self._calculate_experience_years(result['experience'], text)
        
        # Extract projects
        result['projects'] = self._extract_projects(text)
        
        # Extract education (new format: degree, institute, passing_year)
        result['education'] = self._extract_education_new_format(text)
        
        # Detect gender
        result['gender'] = self._detect_gender(result['full_name'], text)
        logger.info(f"🚻 Detected gender: {result['gender']} for name: {result['full_name']}")
        
        return result
    
    def _extract_name(self, text: str, doc) -> Optional[str]:
        """Extract candidate name with improved logic"""
        # Strategy 1: Look at first 5 lines for capitalized names
        lines = text.split('\n')[:10]
        for line in lines:
            line = line.strip()
            # Name is usually 2-4 words, all capitalized or title case
            words = line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
                # Check if it's not a heading
                if not any(keyword in line.lower() for keyword in ['resume', 'cv', 'curriculum', 'vitae']):
                    return line
        
        # Strategy 2: Use spaCy PERSON entities
        persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if persons:
            # Filter out common non-names
            for person in persons:
                if len(person.split()) >= 2 and len(person) < 50:
                    return person
        
        return None
    
    def _detect_gender(self, name: str, text: str) -> Optional[str]:
        """Detect gender from name and resume text"""
        if not name:
            logger.warning("⚠️ Gender detection: No name provided")
            return None
        
        # Extract first name
        first_name = name.split()[0].replace('.', '').strip()
        if not first_name:
            return None
            
        first_name_clean = first_name.title()
        logger.info(f"🔍 Checking first name: '{first_name_clean}' from full name: '{name}'")
        
        # 1. Try Genderize.io API (Removed to save space/deps)
        logger.info(f"🌍 Genderize.io API skipped (removed to save space)")

        # 2. Fallback to Internal Database
        # Common Indian male names (first names)
        male_names = {
            'rahul', 'amit', 'raj', 'vijay', 'ajay', 'anil', 'suresh', 'ramesh', 'mahesh',
            'rakesh', 'deepak', 'sandeep', 'sanjay', 'manoj', 'rajesh', 'ravi', 'ashok',
            'kumar', 'arun', 'ankit', 'rohit', 'nikhil', 'vishal', 'gaurav', 'kuldeep',
            'pradeep', 'naveen', 'karan', 'arjun', 'varun', 'harsh', 'yash', 'dev',
            'aman', 'akash', 'abhishek', 'aditya', 'akhil', 'anuj', 'aryan', 'ayush',
            'mohammed', 'muhammad', 'sai', 'krishna', 'shiva', 'ganesh'
        }
        
        # Common Indian female names (first names)
        female_names = {
            'priya', 'anjali', 'pooja', 'neha', 'kavita', 'sunita', 'rekha', 'geeta',
            'seema', 'ritu', 'nisha', 'asha', 'meera', 'radha', 'sita', 'lata',
            'anita', 'savita', 'shilpa', 'sneha', 'divya', 'swati', 'preeti', 'jyoti',
            'riya', 'sakshi', 'shruti', 'simran', 'tanvi', 'ishita', 'khushi', 'ananya',
            'aarti', 'deepika', 'kritika', 'megha', 'nikita', 'pallavi', 'richa', 'sonal',
            'lakshmi', 'saraswati', 'durga', 'parvati'
        }
        
        first_name_lower = first_name_clean.lower()
        if first_name_lower in male_names:
            logger.info(f"✅ Matched '{first_name_clean}' in internal male names DB")
            return 'Male'
        elif first_name_lower in female_names:
            logger.info(f"✅ Matched '{first_name_clean}' in internal female names DB")
            return 'Female'
        
        # 3. Fallback to Pronoun Analysis
        text_lower = text.lower()
        
        male_indicators = ['he ', 'his ', 'him ', 'mr.', 'mr ', 'himself']
        female_indicators = ['she ', 'her ', 'hers ', 'ms.', 'ms ', 'mrs.', 'mrs ', 'herself', 'miss ']
        
        male_count = sum(text_lower.count(indicator) for indicator in male_indicators)
        female_count = sum(text_lower.count(indicator) for indicator in female_indicators)
        
        # Require significant evidence (count > 2) and clear winner
        if male_count > female_count and male_count > 2:
            logger.info(f"✅ Detected Male based on pronouns (M:{male_count} vs F:{female_count})")
            return 'Male'
        elif female_count > male_count and female_count > 2:
            logger.info(f"✅ Detected Female based on pronouns (F:{female_count} vs M:{male_count})")
            return 'Female'
        
        logger.warning(f"⚠️ Could not determine gender for '{name}'")
        return None
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        matches = self.email_pattern.findall(text)
        # Return first non-noreply email
        for email in matches:
            if 'noreply' not in email.lower():
                return email.lower()
        return matches[0].lower() if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        matches = self.phone_pattern.findall(text)
        if matches:
            # The regex now uses a capturing group, so matches[0] is already the number part
            phone = matches[0].replace(' ', '').replace('-', '').replace('\t', '')
            # Return only the 10-digit number
            if len(phone) == 10:
                return phone
            elif len(phone) > 10:
                return phone[-10:]
            else:
                return None
        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        matches = self.linkedin_pattern.findall(text)
        return f"https://linkedin.com/in/{matches[0]}" if matches else None
    
    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL"""
        matches = self.github_pattern.findall(text)
        return f"https://github.com/{matches[0]}" if matches else None
    
    def _extract_skills(self, text: str) -> List[str]:
        """
        Extract technical and soft skills using regex + keyword matching.
        Enhanced version with comprehensive skill detection.
        """
        text_lower = text.lower()
        found_skills = set()
        
        # Match skills from database with word boundaries
        for skill in self.skills_database:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill.title())
        
        # Extract from Skills section specifically
        skills_section_pattern = r'(?i)(skills?|technical skills?|core competencies|tools & technologies)(.*?)(?=\n\s*\n|experience|education|projects|$)'
        match = re.search(skills_section_pattern, text, re.DOTALL)
        if match:
            skills_text = match.group(2)
            # Find bullet points and comma-separated lists
            for line in skills_text.split('\n'):
                line = line.strip('•●○-– \t')
                if line:
                    # Check each word against database
                    for skill in self.skills_database:
                        if skill.lower() in line.lower():
                            found_skills.add(skill.title())
        
        return sorted(list(found_skills))
    
    def _extract_skills_improved(self, text: str) -> List[str]:
        """Alias for _extract_skills for backward compatibility"""
        return self._extract_skills(text)
    
    def _extract_experience_improved(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience with better parsing"""
        experiences = []
        
        # Find experience section
        exp_pattern = r'(?i)(experience|work history|employment|professional experience)(.*?)(?=education|skills|projects|certifications|$)'
        match = re.search(exp_pattern, text, re.DOTALL)
        
        if not match:
            return []
        
        exp_text = match.group(2)
        
        # Split by date patterns or double newlines
        entries = re.split(r'\n\s*\n+', exp_text)
        
        for entry in entries:
            entry = entry.strip()
            if len(entry) < 30:  # Too short to be meaningful
                continue
                
            experience = {}
            lines = entry.split('\n')
            
            # Extract dates (various formats)
            date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)'
            date_match = re.search(date_pattern, entry, re.I)
            if date_match:
                experience['start_date'] = date_match.group(1)
                experience['end_date'] = date_match.group(2)
            
            # Extract job title (usually first line)
            experience['title'] = lines[0].strip() if lines else None
            
            # Extract company (usually second line or line with company indicators)
            for line in lines[1:3]:
                if any(indicator in line.lower() for indicator in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies']):
                    experience['company'] = line.strip()
                    break
            
            experience['description'] = entry
            experiences.append(experience)
            
            if len(experiences) >= 5:  # Limit to 5 most recent
                break
        
        return experiences
    
    def _extract_designation(self, text: str) -> Optional[str]:
        """Extract current designation/job title"""
        # Look for designation patterns in first few lines
        lines = text.split('\n')[:15]
        
        # Common designation keywords
        designation_keywords = ['engineer', 'developer', 'designer', 'manager', 'analyst', 
                               'consultant', 'architect', 'lead', 'senior', 'junior']
        
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in designation_keywords):
                # Skip if it's a section header
                if not any(header in line_lower for header in ['experience', 'education', 'skills']):
                    return line.strip()
        
        return None
    
    def _extract_location_details(self, text: str) -> Dict[str, Optional[str]]:
        """Extract country and city from resume"""
        result = {'country': None, 'city': None}
        
        # Common Indian cities
        indian_cities = ['mumbai', 'delhi', 'bangalore', 'hyderabad', 'chennai', 'kolkata', 
                        'pune', 'ahmedabad', 'jaipur', 'lucknow', 'kanpur', 'nagpur', 
                        'indore', 'bhopal', 'visakhapatnam', 'patna', 'vadodara', 'ghaziabad',
                        'ludhiana', 'agra', 'nashik', 'faridabad', 'meerut', 'rajkot', 'palwal']
        
        text_lower = text.lower()
        
        # Check for India
        if 'india' in text_lower or any(city in text_lower for city in indian_cities):
            result['country'] = 'India'
        
        # Extract city
        for city in indian_cities:
            if city in text_lower:
                result['city'] = city.title()
                break
        
        return result
    
    def _extract_experience_new_format(self, text: str) -> List[Dict[str, Any]]:
        """Extract experience in new format: company, role, start_year, end_year"""
        experiences = []
        
        # Find experience section
        exp_pattern = r'(?i)(experience|work history|employment|professional experience)(.*?)(?=education|skills|projects|certifications|$)'
        match = re.search(exp_pattern, text, re.DOTALL)
        
        if not match:
            return []
        
        exp_text = match.group(2)
        entries = re.split(r'\n\s*\n+', exp_text)
        
        for entry in entries:
            entry = entry.strip()
            if len(entry) < 30:
                continue
            
            experience = {}
            lines = entry.split('\n')
            
            # Extract role (usually first line)
            experience['role'] = lines[0].strip() if lines else None
            
            # Extract company
            for line in lines[1:3]:
                if any(indicator in line.lower() for indicator in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'animation', 'pvt']):
                    experience['company'] = line.strip()
                    break
            
            # Extract years (2024, 2024-2025, etc.)
            year_pattern = r'\b(20\d{2})\s*[-–—]?\s*(20\d{2}|Present|Current)?\b'
            year_matches = re.findall(year_pattern, entry, re.I)
            if year_matches:
                first_match = year_matches[0]
                experience['start_year'] = int(first_match[0])
                if first_match[1] and first_match[1].isdigit():
                    experience['end_year'] = int(first_match[1])
                elif 'present' in entry.lower() or 'current' in entry.lower():
                    experience['end_year'] = datetime.now().year
            
            if experience.get('company') or experience.get('role'):
                experiences.append(experience)
            
            if len(experiences) >= 5:
                break
        
        return experiences
    
    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract projects with title and description"""
        projects = []
        
        # Find projects section
        proj_pattern = r'(?i)(projects?|portfolio)(.*?)(?=education|experience|skills|certifications|languages|hobbies|$)'
        match = re.search(proj_pattern, text, re.DOTALL)
        
        if not match:
            return []
        
        proj_text = match.group(2)
        entries = re.split(r'\n\s*\n+', proj_text)
        
        for entry in entries:
            entry = entry.strip()
            if len(entry) < 20:
                continue
            
            lines = entry.split('\n')
            project = {
                'project_title': lines[0].strip() if lines else None,
                'description': '\n'.join(lines[1:]).strip() if len(lines) > 1 else entry
            }
            
            if project['project_title']:
                projects.append(project)
            
            if len(projects) >= 5:
                break
        
        return projects
    
    def _calculate_experience_years(self, experiences: List[Dict], text: str) -> float:
        """Calculate total years of experience from new format"""
        if not experiences:
            # Fallback: Look for "X years of experience" in text
            exp_mention = re.search(r'(\d+)[\s+]years?\s+(?:of\s+)?experience', text, re.I)
            if exp_mention:
                return float(exp_mention.group(1))
            return 0.0
        
        total_years = 0.0
        
        for exp in experiences:
            if 'start_year' in exp and 'end_year' in exp:
                years = exp['end_year'] - exp['start_year']
                total_years += max(0, years)
        
        return round(total_years, 1)
    
    def _extract_education_new_format(self, text: str) -> List[Dict[str, Any]]:
        """Extract education with better parsing"""
        education = []
        
        # Find education section
        edu_pattern = r'(?i)(education|academic|qualification|degrees?)(.*?)(?=experience|skills|projects|certifications|$)'
        match = re.search(edu_pattern, text, re.DOTALL)
        
        if not match:
            return []
        
        edu_text = match.group(2)
        entries = re.split(r'\n\s*\n+', edu_text)
        
        degree_keywords = ['bachelor', 'master', 'phd', 'doctorate', 'diploma', 'associate', 'mba', 'b.tech', 'm.tech', 'b.e', 'm.e', 'b.sc', 'm.sc', 'b.com', 'm.com', 'bca', 'mca']
        
        for entry in entries:
            entry = entry.strip()
            if len(entry) < 20:
                continue
            
            edu_entry = {}
            
            # Extract degree - look for degree keywords in the text
            degree_found = None
            for degree in degree_keywords:
                # Match degree with word boundaries
                pattern = r'\b' + re.escape(degree) + r'\b'
                if re.search(pattern, entry.lower()):
                    # Try to extract the full degree line
                    lines = entry.split('\n')
                    for line in lines:
                        if degree in line.lower():
                            degree_found = line.strip()
                            break
                    if not degree_found:
                        degree_found = degree.upper()
                    break
            
            if degree_found:
                edu_entry['degree'] = degree_found
            
            # Extract institute name (renamed from institution_name)
            # Look for university, college, institute, school keywords
            institution_pattern = r'([A-Z][A-Za-z\s,\.&]+(?:University|College|Institute|School|Academy|Board)[A-Za-z\s,\.]*)'
            inst_match = re.search(institution_pattern, entry)
            if inst_match:
                edu_entry['institute'] = inst_match.group(1).strip()
            
            # Extract passing_year (renamed from year)
            # Handle formats like: 2024, 2024-2025, (2024–2025), 2024 - 2025
            year_pattern = r'\b(20\d{2})(?:\s*[-–—]\s*(20\d{2}))?\b'
            year_matches = re.findall(year_pattern, entry)
            if year_matches:
                # Get the last year (graduation year)
                last_match = year_matches[-1]
                edu_entry['passing_year'] = int(last_match[1]) if last_match[1] else int(last_match[0])
            
            if edu_entry.get('degree') or edu_entry.get('institute'):
                education.append(edu_entry)
            
            if len(education) >= 3:
                break
        
        return education
    
    def _calculate_experience(self, experiences: List[Dict], text: str) -> float:
        """Calculate total years of experience"""
        if not experiences:
            # Fallback: Look for "X years of experience" in text
            exp_mention = re.search(r'(\d+)[\s+]years?\s+(?:of\s+)?experience', text, re.I)
            if exp_mention:
                return float(exp_mention.group(1))
            return 0.0
        
        total_years = 0.0
        
        for exp in experiences:
            if 'start_date' in exp and 'end_date' in exp:
                try:
                    # Extract years from dates
                    start_year = int(re.search(r'\d{4}', exp['start_date']).group())
                    
                    end_str = exp['end_date'].lower()
                    if 'present' in end_str or 'current' in end_str:
                        end_year = datetime.now().year
                    else:
                        end_year = int(re.search(r'\d{4}', exp['end_date']).group())
                    
                    years = end_year - start_year
                    total_years += max(0, years)
                except:
                    pass
        
        return round(total_years, 1)
